"""Tests del reemplazo dentro del .docx.

Esta es la parte del proyecto que puede fallar en silencio: si el reemplazo no
encuentra el texto, entregas un contrato con los datos del cliente anterior y
nadie se entera hasta que alguien lo firma. Por eso es la que más se testea.
"""

from docx import Document

from contratos.reemplazo import Reemplazo, aplicar


def _texto_completo(documento) -> str:
    partes = [p.text for p in documento.paragraphs]
    for tabla in documento.tables:
        for fila in tabla.rows:
            partes.extend(celda.text for celda in fila.cells)
    return "\n".join(partes)


class TestReemplazoSimple:
    def test_cambia_el_texto_de_un_parrafo(self):
        doc = Document()
        doc.add_paragraph("El Cliente ACME S.A. declara.")

        aplicar(doc, [Reemplazo("cliente", "ACME S.A.", "Beta Ltda.")])

        assert "Beta Ltda." in _texto_completo(doc)
        assert "ACME" not in _texto_completo(doc)

    def test_cuenta_cada_aparicion(self):
        doc = Document()
        doc.add_paragraph("ACME S.A. y ACME S.A. otra vez.")
        doc.add_paragraph("Y aquí ACME S.A. de nuevo.")

        reemplazo = Reemplazo("cliente", "ACME S.A.", "Beta Ltda.")
        aplicar(doc, [reemplazo])

        assert reemplazo.veces == 3

    def test_reporta_cuando_no_encuentra_nada(self):
        doc = Document()
        doc.add_paragraph("Texto sin relación.")

        reemplazo = Reemplazo("cliente", "ACME S.A.", "Beta Ltda.")
        aplicar(doc, [reemplazo])

        assert reemplazo.veces == 0
        assert not reemplazo.aplicado


class TestRunsPartidas:
    """El caso que rompe cualquier implementación ingenua."""

    def test_encuentra_texto_partido_entre_varias_runs(self):
        doc = Document()
        parrafo = doc.add_paragraph()
        parrafo.add_run("Gran Amér")
        parrafo.add_run("icas Santiago")   # Word parte así todo el tiempo
        parrafo.add_run(" Chile S.A.")

        reemplazo = Reemplazo(
            "cliente", "Gran Américas Santiago Chile S.A.", "Transportes Beta Ltda."
        )
        aplicar(doc, [reemplazo])

        assert reemplazo.veces == 1
        assert doc.paragraphs[0].text == "Transportes Beta Ltda."

    def test_conserva_el_texto_que_rodea_al_hallazgo(self):
        doc = Document()
        parrafo = doc.add_paragraph()
        parrafo.add_run("Entre ACM")
        parrafo.add_run("E S.A. y Volvo Chile.")

        aplicar(doc, [Reemplazo("cliente", "ACME S.A.", "Beta")])

        assert doc.paragraphs[0].text == "Entre Beta y Volvo Chile."


class TestOrdenYColisiones:
    def test_el_valor_largo_se_procesa_antes_que_el_corto(self):
        """La trampa de 'Santiago'.

        Si se reemplazara la dirección antes que la razón social, el nombre
        'Gran Américas Santiago Chile S.A.' quedaría mutilado.
        """
        doc = Document()
        doc.add_paragraph("Gran Américas Santiago Chile S.A., ciudad de Santiago.")

        nombre = Reemplazo("cliente", "Gran Américas Santiago Chile S.A.", "Beta Ltda.")
        ciudad = Reemplazo("direccion_cliente", "ciudad de Santiago", "ciudad de Concepción")
        aplicar(doc, [ciudad, nombre])  # a propósito en el orden "malo"

        assert doc.paragraphs[0].text == "Beta Ltda., ciudad de Concepción."
        assert nombre.veces == 1
        assert ciudad.veces == 1

    def test_no_entra_en_bucle_si_el_valor_nuevo_contiene_al_viejo(self):
        doc = Document()
        doc.add_paragraph("Volvo firma.")

        reemplazo = Reemplazo("sublicenciante", "Volvo", "Volvo Chile SpA")
        aplicar(doc, [reemplazo])

        assert reemplazo.veces == 1
        assert doc.paragraphs[0].text == "Volvo Chile SpA firma."

    def test_ignora_los_campos_sin_cambio(self):
        doc = Document()
        doc.add_paragraph("ACME S.A.")

        reemplazo = Reemplazo("cliente", "ACME S.A.", "ACME S.A.")
        aplicar(doc, [reemplazo])

        assert reemplazo.veces == 0


class TestOtrosContenedores:
    def test_tambien_reemplaza_dentro_de_tablas(self):
        doc = Document()
        tabla = doc.add_table(rows=1, cols=2)
        tabla.rows[0].cells[0].text = "Cliente"
        tabla.rows[0].cells[1].text = "ACME S.A."

        reemplazo = Reemplazo("cliente", "ACME S.A.", "Beta Ltda.")
        aplicar(doc, [reemplazo])

        assert reemplazo.veces == 1
        assert tabla.rows[0].cells[1].text == "Beta Ltda."

    def test_tambien_reemplaza_en_el_encabezado(self):
        doc = Document()
        doc.sections[0].header.paragraphs[0].text = "Contrato ACME S.A."

        reemplazo = Reemplazo("cliente", "ACME S.A.", "Beta Ltda.")
        aplicar(doc, [reemplazo])

        assert reemplazo.veces == 1
