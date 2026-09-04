"""Tests sobre el contenido de las plantillas .docx oficiales.

No prueban codigo: prueban que los documentos que se entregan al cliente
digan lo que tienen que decir. Es contenido legal, y un cambio silencioso
ahi no lo atrapa ningun otro test.
"""

from datetime import date

import pytest
from docx import Document

from contratos.modelos import Cliente, ContratoImpactOnline, ContratoTechTool
from contratos.motor import generar_docx

REPRESENTANTE = "Rodrigo Nunez Salas"

#: Nombres que estaban escritos a firme en la pagina de firmas y que ya no
#: deben aparecer: el cliente firma con su propio representante legal.
NOMBRES_RETIRADOS = ("Fuentealba", "Juan Nicolás Palma", "Juan Nicolas Palma")


@pytest.fixture
def cliente():
    return Cliente(
        razon_social="Transportes Del Sur Ltda.",
        rut="77981536-6",
        direccion="Avenida Siempre Viva 742, Santiago",
        contacto="Paula Contreras",
        representante_legal=REPRESENTANTE,
    )


def _texto_completo(ruta) -> str:
    doc = Document(ruta)
    partes = [p.text for p in doc.paragraphs]
    partes += [c.text for t in doc.tables for f in t.rows for c in f.cells]
    return "\n".join(partes)


def _tabla_de_firmas(ruta):
    doc = Document(ruta)
    for tabla in doc.tables:
        texto = "\n".join(c.text for f in tabla.rows for c in f.cells)
        if "Nombre completo" in texto and "CLIENTE" in texto:
            return tabla
    raise AssertionError("La plantilla no tiene tabla de firmas.")


@pytest.mark.parametrize("cliente_externo", [False, True])
class TestFirmasImpact:
    def _generar(self, cliente, cliente_externo, tmp_path):
        contrato = ContratoImpactOnline(
            cliente=cliente,
            fecha=date(2026, 9, 4),
            cliente_id="CL123456",
            cliente_externo=cliente_externo,
        )
        return generar_docx(contrato, dir_salida=tmp_path)

    def test_el_cliente_firma_con_su_representante_legal(
        self, cliente, cliente_externo, tmp_path
    ):
        ruta = self._generar(cliente, cliente_externo, tmp_path)
        tabla = _tabla_de_firmas(ruta)

        # La seccion CLIENTE va antes de la de VOLVO: se corta ahi.
        filas = ["\n".join(c.text for c in f.cells) for f in tabla.rows]
        corte = next(i for i, f in enumerate(filas) if "VOLVO CHILE" in f)
        seccion_cliente = "\n".join(filas[:corte])

        assert REPRESENTANTE in seccion_cliente

    def test_no_quedan_los_nombres_fijos_antiguos(
        self, cliente, cliente_externo, tmp_path
    ):
        texto = _texto_completo(self._generar(cliente, cliente_externo, tmp_path))
        for nombre in NOMBRES_RETIRADOS:
            assert nombre not in texto

    def test_hay_una_sola_firma_de_cliente(self, cliente, cliente_externo, tmp_path):
        ruta = self._generar(cliente, cliente_externo, tmp_path)
        tabla = _tabla_de_firmas(ruta)

        filas = ["\n".join(c.text for c in f.cells) for f in tabla.rows]
        corte = next(i for i, f in enumerate(filas) if "VOLVO CHILE" in f)
        bloques = [f for f in filas[:corte] if "Nombre completo" in f]

        assert len(bloques) == 1

    def test_la_firma_de_volvo_sigue_fija(self, cliente, cliente_externo, tmp_path):
        texto = _texto_completo(self._generar(cliente, cliente_externo, tmp_path))
        assert "Marcela Ugalde Flaquer" in texto


class TestFirmasTechTool:
    @pytest.mark.parametrize("cliente_externo", [False, True])
    def test_no_arrastra_nombres_de_firmantes_fijos(
        self, cliente, cliente_externo, tmp_path
    ):
        contrato = ContratoTechTool(
            cliente=cliente,
            fecha=date(2026, 9, 4),
            partner_id="CL999888",
            cliente_externo=cliente_externo,
        )
        texto = _texto_completo(generar_docx(contrato, dir_salida=tmp_path))
        for nombre in NOMBRES_RETIRADOS:
            assert nombre not in texto
