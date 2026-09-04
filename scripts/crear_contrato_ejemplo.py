"""Crea un contrato de ejemplo YA LLENADO, para probar la subida.

No lleva marcadores {{ }}: es un contrato normal, como los que vas a subir.

Uso:  python scripts/crear_contrato_ejemplo.py
"""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

DESTINO = Path(__file__).resolve().parents[1] / "ejemplos" / "contrato-ejemplo.docx"

CLIENTE = [
    ("Cliente", "Gran Américas Santiago Chile S.A."),
    ("RUT", "77.981.536-6"),
    ("Dirección", "Eliodoro Yáñez N° 2972, oficina 311, comuna de Providencia, ciudad de Santiago"),
    ("Software", "Impact Online"),
    ("Cliente ID", "CL099033"),
    ("Digipass (es)", "NA"),
    ("Contacto", "Katherine Casas"),
    ("Representante legal", "Claudio Castillo Castillo"),
]

SUBLICENCIANTE = [
    ("Empresa del Grupo Volvo", "VOLVO CHILE SPA."),
    ("RUT", "76.284.920-8"),
    ("Dirección", "Avenida Presidente Eduardo Frei Montalva 8691, Quilicura, Santiago"),
]

CLAUSULAS = [
    ("1. LICENCIA DE USO",
     "El Sublicenciante concede a Gran Américas Santiago Chile S.A. el derecho no "
     "exclusivo de uso del Software Impact Online. El Software se debe usar por el "
     "Cliente para informaciones sobre repuestos y servicios."),
    ("2. CONTRAPRESTACIÓN",
     "La sublicencia de uso se concede sin costo para el Cliente, asumiendo VOLVO "
     "CHILE SPA. íntegramente la tarifa asociada al servicio, cuyo valor referencial "
     "asciende a 3,7 UF mensuales."),
    ("3. DOMICILIO Y JURISDICCIÓN",
     "Para todos los efectos del presente Contrato las partes fijan su domicilio en "
     "la ciudad y comuna de Santiago y se someten a la jurisdicción de los tribunales "
     "ordinarios de justicia con asiento en dicha comuna."),
]


def tabla_de(doc, filas):
    tabla = doc.add_table(rows=0, cols=2)
    tabla.style = "Table Grid"
    for etiqueta, valor in filas:
        celdas = tabla.add_row().cells
        celdas[0].text = etiqueta
        celdas[1].text = valor
    return tabla


def construir():
    doc = Document()
    doc.styles["Normal"].font.size = Pt(10)

    titulo = doc.add_paragraph()
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    corrida = titulo.add_run("CONTRATO DE SUBLICENCIA DEL SOFTWARE (Impact Online)")
    corrida.bold = True
    corrida.font.size = Pt(13)

    doc.add_paragraph('Este Contrato se celebra con fecha 22-07-2026 entre ("Sublicenciatario"):')
    tabla_de(doc, CLIENTE)

    doc.add_paragraph()
    doc.add_paragraph('y ("Sublicenciante"):')
    tabla_de(doc, SUBLICENCIANTE)

    doc.add_paragraph()
    for encabezado, cuerpo in CLAUSULAS:
        parrafo = doc.add_paragraph()
        parrafo.add_run(encabezado).bold = True
        doc.add_paragraph(cuerpo)

    DESTINO.parent.mkdir(parents=True, exist_ok=True)
    doc.save(DESTINO)
    print(f"Contrato de ejemplo creado en: {DESTINO}")


if __name__ == "__main__":
    construir()
