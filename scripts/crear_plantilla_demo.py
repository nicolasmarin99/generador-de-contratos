"""Genera una plantilla .docx de ejemplo con marcadores {{ }}.

Solo sirve para probar el proyecto sin tener las plantillas reales.
Cuando marques tus propios contratos, sobreescribe plantillas/impact_online.docx
y este script deja de ser necesario.

Uso:  python scripts/crear_plantilla_demo.py
"""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

RAIZ = Path(__file__).resolve().parents[1]
DESTINO = RAIZ / "plantillas" / "impact_online.docx"

CAMPOS = [
    ("Cliente", "{{ cliente }}"),
    ("RUT", "{{ rut_cliente }}"),
    ("Direccion", "{{ direccion_cliente }}"),
    ("Software", "{{ software }}"),
    ("Cliente ID", "{{ cliente_id }}"),
    ("Digipass (es)", "{{ digipass }}"),
    ("Contacto", "{{ contacto }}"),
    ("Representante legal", "{{ representante_legal }}"),
]

CLAUSULA = (
    "El Sublicenciante concede al Cliente el derecho no exclusivo de uso del "
    "Software Impact Online. El Software se debe usar por el Cliente para "
    "informaciones sobre repuestos y servicios. El Cliente esta de acuerdo que "
    "no podra sublicenciar el Software, salvo autorizacion expresa y por escrito "
    "del Sublicenciante."
)


def construir() -> None:
    doc = Document()
    doc.styles["Normal"].font.size = Pt(10)

    titulo = doc.add_paragraph()
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    corrida = titulo.add_run("CONTRATO DE SUBLICENCIA DEL SOFTWARE")
    corrida.bold = True
    corrida.font.size = Pt(14)

    doc.add_paragraph('Este Contrato se celebra con fecha {{ fecha }} entre ("Sublicenciatario"):')

    tabla = doc.add_table(rows=0, cols=2)
    tabla.style = "Table Grid"
    for etiqueta, marcador in CAMPOS:
        fila = tabla.add_row().cells
        fila[0].text = etiqueta
        fila[1].text = marcador

    doc.add_paragraph()
    doc.add_paragraph('y ("Sublicenciante"):')

    tabla2 = doc.add_table(rows=0, cols=2)
    tabla2.style = "Table Grid"
    for etiqueta, marcador in [
        ("Empresa del Grupo Volvo", "{{ sublicenciante }}"),
        ("RUT", "{{ rut_sublicenciante }}"),
        ("Direccion", "{{ direccion_sublicenciante }}"),
    ]:
        fila = tabla2.add_row().cells
        fila[0].text = etiqueta
        fila[1].text = marcador

    doc.add_paragraph()
    encabezado = doc.add_paragraph()
    encabezado.add_run("1. LICENCIA DE USO").bold = True
    doc.add_paragraph(CLAUSULA)

    DESTINO.parent.mkdir(parents=True, exist_ok=True)
    doc.save(DESTINO)
    print(f"Plantilla de demostracion creada en: {DESTINO}")


if __name__ == "__main__":
    construir()
