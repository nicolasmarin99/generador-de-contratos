"""Reemplazo de texto dentro de un .docx sin romper el formato.

EL PROBLEMA
-----------
Word no guarda un parrafo como una cadena continua. Lo parte en fragmentos
llamados "runs" cada vez que cambia el formato o el corrector deja una marca.
Un nombre que en pantalla se ve entero puede estar guardado asi:

    <w:r><w:t>Gran Amér</w:t></w:r><w:r><w:t>icas Santiago</w:t></w:r>

Por eso un `parrafo.text.replace(...)` ingenuo no encuentra nada: el texto que
buscas no existe como cadena contigua en ninguna run.

LA SOLUCION
-----------
1. Concatenar el texto de todas las runs del parrafo.
2. Buscar sobre esa cadena completa.
3. Al encontrar una coincidencia, escribir el reemplazo en la run donde
   empieza y borrar los caracteres sobrantes de las runs siguientes.

Asi el formato de cada run se conserva, salvo el de la porcion reemplazada,
que hereda el de la primera run del hallazgo. Para nombres, RUT y direcciones
eso es exactamente lo que quieres.
"""

from dataclasses import dataclass
from typing import Iterator

from docx.document import Document as DocumentoWord
from docx.table import Table
from docx.text.paragraph import Paragraph


@dataclass
class Reemplazo:
    """Una sustitucion pedida por el usuario y su resultado."""

    clave: str
    valor_anterior: str
    valor_nuevo: str
    veces: int = 0

    @property
    def aplicado(self) -> bool:
        return self.veces > 0

    @property
    def sin_cambios(self) -> bool:
        return self.valor_anterior == self.valor_nuevo


def aplicar(documento: DocumentoWord, reemplazos: list[Reemplazo]) -> list[Reemplazo]:
    """Aplica los reemplazos sobre todo el documento.

    Devuelve la misma lista con el conteo `veces` actualizado, para que quien
    llame pueda mostrarle al usuario que se cambio y que no. Ese reporte es la
    red de seguridad del sistema: sin el, un reemplazo que no encontro nada
    pasaria inadvertido y entregarias un contrato con datos del cliente viejo.
    """
    activos = [
        r for r in reemplazos
        if r.valor_anterior and not r.sin_cambios
    ]

    # ORDEN CRITICO: de mas largo a mas corto.
    #
    # Si "Santiago" se reemplazara antes que "Gran Américas Santiago Chile S.A.",
    # el nombre largo quedaria mutilado y ya no coincidiria con nada. Procesar
    # primero las cadenas largas evita que las cortas les coman pedazos.
    activos.sort(key=lambda r: len(r.valor_anterior), reverse=True)

    for parrafo in _todos_los_parrafos(documento):
        _reemplazar_en_parrafo(parrafo, activos)

    return reemplazos


def _reemplazar_en_parrafo(parrafo: Paragraph, reemplazos: list[Reemplazo]) -> None:
    runs = parrafo.runs
    if not runs:
        return

    textos = [run.text for run in runs]
    if not any(textos):
        return

    hubo_cambio = False

    for reemplazo in reemplazos:
        desde = 0
        while True:
            completo = "".join(textos)
            posicion = completo.find(reemplazo.valor_anterior, desde)
            if posicion == -1:
                break

            _sustituir(textos, posicion, len(reemplazo.valor_anterior), reemplazo.valor_nuevo)
            reemplazo.veces += 1
            hubo_cambio = True

            # Continuar despues del texto ya insertado. Evita un bucle infinito
            # cuando el valor nuevo contiene al viejo (ej: "Volvo" -> "Volvo Chile").
            desde = posicion + len(reemplazo.valor_nuevo)

    if hubo_cambio:
        for run, texto in zip(runs, textos):
            run.text = texto


def _sustituir(textos: list[str], inicio: int, largo: int, nuevo: str) -> None:
    """Reemplaza un tramo de la cadena virtual, repartido entre varias runs.

    `textos` se modifica en el sitio. El texto nuevo entra completo en la
    primera run tocada; de las demas solo se borran los caracteres del tramo.
    """
    fin = inicio + largo
    posicion = 0
    ya_insertado = False

    for indice, texto in enumerate(textos):
        inicio_run = posicion
        fin_run = posicion + len(texto)
        posicion = fin_run  # se avanza con la longitud ORIGINAL, no la nueva

        if fin_run <= inicio or inicio_run >= fin:
            continue  # esta run queda fuera del tramo

        recorte_desde = max(inicio, inicio_run) - inicio_run
        recorte_hasta = min(fin, fin_run) - inicio_run

        insercion = "" if ya_insertado else nuevo
        ya_insertado = True

        textos[indice] = texto[:recorte_desde] + insercion + texto[recorte_hasta:]


def _todos_los_parrafos(documento: DocumentoWord) -> Iterator[Paragraph]:
    """Recorre TODO el documento, no solo el cuerpo principal.

    Los datos del cliente pueden estar en tablas, encabezados o pies de pagina.
    Olvidar alguno de esos contenedores es el segundo error mas comun despues
    del de las runs.
    """
    yield from documento.paragraphs

    for tabla in documento.tables:
        yield from _parrafos_de_tabla(tabla)

    for seccion in documento.sections:
        for contenedor in (
            seccion.header,
            seccion.footer,
            seccion.first_page_header,
            seccion.first_page_footer,
            seccion.even_page_header,
            seccion.even_page_footer,
        ):
            if contenedor is None:
                continue
            yield from contenedor.paragraphs
            for tabla in contenedor.tables:
                yield from _parrafos_de_tabla(tabla)


def _parrafos_de_tabla(tabla: Table) -> Iterator[Paragraph]:
    for fila in tabla.rows:
        for celda in fila.cells:
            yield from celda.paragraphs
            for anidada in celda.tables:  # las tablas pueden contener tablas
                yield from _parrafos_de_tabla(anidada)
